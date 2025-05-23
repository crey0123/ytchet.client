from flask import Flask, render_template, request, redirect, url_for, flash, session
from flask_bcrypt import Bcrypt
import MySQLdb
import json
import openpyxl  # Для работы с Excel
from io import BytesIO  # Для создания временного файла в памяти
from flask import Response  # Для возврата Excel-файла пользователю
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.secret_key = 'your_secret_key'
bcrypt = Bcrypt(app)

# Подключение к базе данных
db = MySQLdb.connect(
    host="localhost",
    user="root",
    passwd="Ilya20514",
    db="saqt",
    charset="utf8mb4"
)

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        # Проверка пользователя в базе данных
        cursor = db.cursor()
        cursor.execute("SELECT AdministratorID, Password FROM Administrators WHERE Username = %s", (username,))
        user = cursor.fetchone()

        if user and bcrypt.check_password_hash(user[1], password):
            session["user_id"] = user[0]
            session["username"] = username
            return redirect(url_for("dashboard"))
        else:
            flash("Неверное имя пользователя или пароль.", "danger")
    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    flash("Вы успешно вышли из системы.", "success")
    return redirect(url_for("login"))


@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        flash("Пожалуйста, войдите в систему.", "warning")
        return redirect(url_for("login"))
    return render_template("dashboard.html", username=session["username"])


@app.route("/add_user", methods=["GET", "POST"])
def add_user():
    if "user_id" not in session:
        flash("Пожалуйста, войдите в систему.", "warning")
        return redirect(url_for("login"))

    cursor = db.cursor()

    if request.method == "POST":
        # Получаем данные из формы
        fio = request.form["fio"]
        position_id = request.form["position_id"]
        department_id = request.form["department_id"]
        selected_bases = request.form.getlist("access_bases")
        
        # Обработка дополнительных полей: собираем все ключи, начинающиеся с "additional_fields["
        raw_form = request.form.to_dict(flat=False)
        additional_fields = {}
        for key, value in raw_form.items():
            if key.startswith("additional_fields[") and key.endswith("]"):
                field_name = key[len("additional_fields["):-1]
                additional_fields[field_name] = value[0] if value else None

        # Проверяем заполненность обязательных полей
        if not fio or not position_id or not department_id:
            flash("Все обязательные поля должны быть заполнены.", "danger")
            return redirect(url_for("add_user"))

        try:
            # Добавляем пользователя в таблицу Clients
            cursor.execute("""
                INSERT INTO Clients (fio, PositionID, DepartmentID) 
                VALUES (%s, %s, %s)
            """, (fio, position_id, department_id))
            client_id = cursor.lastrowid

            # Добавляем записи в таблицу Client_AccessBase
            for base_id in selected_bases:
                cursor.execute("""
                    INSERT INTO Client_AccessBase (Client_id, AccessBaseID)
                    VALUES (%s, %s)
                """, (client_id, base_id))

            # Добавляем дополнительные поля в таблицу additional_info
            cursor.execute("""
                INSERT INTO additional_info 
                (Client_id, birthdate, id_nomer, email, phone, profile, ecp_key)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (
                client_id,
                additional_fields.get("birthdate"),
                additional_fields.get("id_nomer"),
                additional_fields.get("email"),
                additional_fields.get("phone"),
                additional_fields.get("profile"),
                additional_fields.get("ecp_key")
            ))

            db.commit()
            flash("Пользователь успешно добавлен.", "success")
        except Exception as e:
            db.rollback()
            flash(f"Ошибка при добавлении пользователя: {str(e)}", "danger")
        finally:
            cursor.close()
        return redirect(url_for("add_user"))

    # GET-запрос: загружаем данные для формы
    try:
        cursor.execute("""
            SELECT PairID
            FROM UserPairs
            WHERE AdministratorID = %s
        """, (session["user_id"],))
        pair_id = cursor.fetchone()

        if pair_id:
            pair_id = pair_id[0]
            cursor.execute("""
                SELECT DepartmentID, DepartmentName
                FROM Departments
                WHERE PairID = %s
            """, (pair_id,))
            departments = cursor.fetchall()
        else:
            departments = []

        cursor.execute("SELECT PositionID, PositionName FROM Positions")
        positions = cursor.fetchall()

        cursor.execute("SELECT AccessBaseID, AccessBaseName FROM AccessBases")
        access_bases = cursor.fetchall()

    except Exception as e:
        flash(f"Ошибка при загрузке данных для формы: {str(e)}", "danger")
        departments, positions, access_bases = [], [], []
    finally:
        cursor.close()

    return render_template("add_user.html", positions=positions, departments=departments, access_bases=access_bases)



@app.route("/get_additional_fields/<int:base_id>", methods=["GET"])
def get_additional_fields(base_id):
    cursor = db.cursor()
    try:
        cursor.execute("SELECT RequiredFields FROM AccessBases WHERE AccessBaseID = %s", (base_id,))
        result = cursor.fetchone()
    except Exception as e:
        return f"Ошибка: {str(e)}", 500
    finally:
        cursor.close()

    if result and result[0]:
        return result[0]  # JSON-формат
    return "[]"


@app.route("/add_position", methods=["POST"])
def add_position():
    if "user_id" not in session:
        return {"status": "error", "message": "Вы не авторизованы"}, 403

    position_name = request.form.get("position_name")
    if not position_name:
        return {"status": "error", "message": "Название должности не может быть пустым"}, 400

    cursor = db.cursor()
    try:
        cursor.execute("INSERT INTO Positions (PositionName) VALUES (%s)", (position_name,))
        db.commit()
        return {"status": "success", "message": "Должность добавлена", "position_id": cursor.lastrowid}
    except Exception as e:
        db.rollback()
        return {"status": "error", "message": str(e)}, 500
    finally:
        cursor.close()





@app.route("/users", methods=["GET"])
def users():
    if "user_id" not in session:
        flash("Пожалуйста, войдите в систему.", "warning")
        return redirect(url_for("login"))

    cursor = db.cursor()
    
    # Получение параметров фильтрации
    department_id = request.args.get("department_id")
    access_base_id = request.args.get("access_base_id")

    query = """
        SELECT c.id, c.fio, p.PositionName, d.DepartmentName, GROUP_CONCAT(ab.AccessBaseName), c.date_added
        FROM Clients c
        LEFT JOIN Positions p ON c.PositionID = p.PositionID
        LEFT JOIN Departments d ON c.DepartmentID = d.DepartmentID
        LEFT JOIN Client_AccessBase cab ON c.id = cab.Client_id
        LEFT JOIN AccessBases ab ON cab.AccessBaseID = ab.AccessBaseID
        WHERE c.deleted_at IS NULL
    """
    filters = []

    # Добавляем фильтр по подразделению
    if department_id:
        query += " AND c.DepartmentID = %s"
        filters.append(department_id)

    # Добавляем фильтр по базе доступа
    if access_base_id:
        query += """
            AND EXISTS (
                SELECT 1 FROM Client_AccessBase cab2 
                WHERE cab2.Client_id = c.id AND cab2.AccessBaseID = %s
            )
        """
        filters.append(access_base_id)

    query += " GROUP BY c.id"

    try:
        cursor.execute(query, filters)
        users = cursor.fetchall()
    except Exception as e:
        flash(f"Ошибка загрузки списка пользователей: {str(e)}", "danger")
        users = []
    finally:
        cursor.close()

    # Загружаем данные для фильтров
    cursor = db.cursor()
    cursor.execute("SELECT DepartmentID, DepartmentName FROM Departments")
    departments = cursor.fetchall()
    cursor.execute("SELECT AccessBaseID, AccessBaseName FROM AccessBases")
    access_bases = cursor.fetchall()
    cursor.close()

    return render_template("users.html", users=users, departments=departments, access_bases=access_bases)


@app.route("/delete_user", methods=["POST", "GET"])
def delete_user():
    if request.method == "GET":
        return render_template("delete_user.html")

    fio = request.form.get("fio")  # Получаем ФИО из формы
    if not fio:
        flash("Пожалуйста, введите ФИО для удаления.", "danger")
        return redirect(url_for("delete_user"))

    cursor = db.cursor()
    try:
        # Проверяем, существует ли пользователь
        cursor.execute("SELECT id FROM Clients WHERE fio = %s", (fio,))
        user = cursor.fetchone()

        if user:
            user_id = user[0]

            # Удаляем зависимые данные перед удалением пользователя
            cursor.execute("DELETE FROM additional_info WHERE Client_id = %s", (user_id,))
            cursor.execute("DELETE FROM Client_AccessBase WHERE Client_id = %s", (user_id,))
            cursor.execute("DELETE FROM Reports WHERE AccessBaseID IN (SELECT AccessBaseID FROM Client_AccessBase WHERE Client_id = %s)", (user_id,))

            # Удаляем самого пользователя
            cursor.execute("DELETE FROM Clients WHERE id = %s", (user_id,))
            db.commit()

            flash(f"Пользователь '{fio}' и вся его информация успешно удалены.", "success")
        else:
            flash(f"Пользователь с ФИО '{fio}' не найден.", "warning")
    except Exception as e:
        db.rollback()
        flash(f"Ошибка при удалении пользователя: {str(e)}", "danger")
    finally:
        cursor.close()

    return redirect(url_for("delete_user"))





def generate_word_report(report_data, headers, base_name):
    """
    Генерирует Word‑документ с текстом, расположенным следующим образом:
    • В верхней части страницы справа – статичный текст договора.
    • Над таблицей слева – динамическая надпись с названием базы.
    • Затем следует таблица с данными отчёта.
    """
    document = Document()
    
    
    # --- Блок 1: Статичный текст договора (выравнивание по правому краю) ---
    contract_text = (
        "Приложение __\n"
        "К Договору № ______ от «__» ______ 20__ г.\n"
        "Между Министерством внутренних дел Республики Беларусь и центральным аппаратом "
        "Следственного комитета Республики Беларусь"
    )
    p_contract = document.add_paragraph(contract_text)
    p_contract.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_contract.runs:
        run.font.size = Pt(11)
    
    # Добавляем небольшой отступ (если требуется)
    document.add_paragraph("")
    
    # --- Блок 2: Динамическая надпись с названием базы (выравнивание по левому краю) ---
    base_info_text = f"Список пользователей, которым продлевается доступ к базе: {base_name}"
    p_base = document.add_paragraph(base_info_text)
    p_base.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_base.runs:
        run.font.size = Pt(11)
    
    # Еще один отступ между текстом и таблицей
    document.add_paragraph("")
    
    # --- Таблица с данными отчёта ---
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Первая строка – заголовки
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Заполнение таблицы данными
    for row in report_data:
        new_row = table.add_row().cells
        for i, cell in enumerate(row):
            new_row[i].text = str(cell) if cell is not None else "Нет данных"
            for paragraph in new_row[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    # Сохраняем документ в BytesIO для дальнейшей отправки по HTTP
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output

@app.route('/generate_report', methods=["GET", "POST"])
def generate_report():
    cursor = db.cursor()

    if request.method == "POST":
        access_base_id = request.form.get("access_base_id")
        start_date = request.form.get("start_date")
        end_date = request.form.get("end_date")
        export_option = request.form.get("export_option", "")  # Значение "excel" или "word"

        if not access_base_id:
            flash("Пожалуйста, выберите базу доступа.", "warning")
            return redirect(url_for("generate_report"))
        
        # Если даты не заданы, используем дату последнего отчёта
        if not start_date or not end_date:
            cursor.execute("SELECT MAX(ReportDate) FROM Reports WHERE AccessBaseID = %s", (access_base_id,))
            last_report_date = cursor.fetchone()[0] or '1900-01-01'
            start_date = start_date or last_report_date
            end_date = end_date or "CURRENT_DATE"
        
        # Получаем список дополнительных полей для выбранной базы
        cursor.execute("SELECT RequiredFields FROM AccessBases WHERE AccessBaseID = %s", (access_base_id,))
        rf_result = cursor.fetchone()
        required_fields = json.loads(rf_result[0]) if rf_result and rf_result[0] else []
        
        # Маппинг дополнительных полей — можно дополнить при необходимости
        field_mapping = {
            "birthdate": "ai.birthdate",
            "id_nomer": "ai.id_nomer",
            "email": "ai.email",
            "phone": "ai.phone",
            "profile": "ai.profile",
            "ecp_key": "ai.ecp_key"
        }
        selected_fields = [field_mapping[field["name"]] for field in required_fields if field["name"] in field_mapping]
        additional_info_query = ", ".join(selected_fields) if selected_fields else "NULL AS no_extra_fields"
        
        # Формируем SQL-запрос для получения данных отчёта
        query = f"""
            SELECT c.fio, p.PositionName, d.DepartmentName, a.AccessBaseName,
                   c.date_added, c.deleted_at, {additional_info_query}
            FROM Clients c
            LEFT JOIN Positions p ON c.PositionID = p.PositionID
            LEFT JOIN Departments d ON c.DepartmentID = d.DepartmentID
            LEFT JOIN Client_AccessBase cab ON c.id = cab.Client_id
            LEFT JOIN AccessBases a ON cab.AccessBaseID = a.AccessBaseID
            LEFT JOIN additional_info ai ON c.id = ai.Client_id
            WHERE cab.AccessBaseID = %s AND c.date_added BETWEEN %s AND %s
        """
        cursor.execute(query, (access_base_id, start_date, end_date))
        report_data = cursor.fetchall()

        # Записываем информацию о сформированном отчёте в базу
        cursor.execute(
            "INSERT INTO Reports (ReportDate, RecordCount, AccessBaseID) VALUES (CURRENT_DATE, %s, %s)",
            (len(report_data), access_base_id)
        )
        db.commit()
        
        # Получаем название базы из данных (предполагается, что AccessBaseName находится в 4-м столбце, индекс 3)
        base_name = "Неизвестная база"
        if report_data and len(report_data) > 0:
            base_name = report_data[0][3]
        
        # Формируем заголовки для отчёта: основные и дополнительные поля
        basic_headers = ["ФИО", "Должность", "Подразделение", "База доступа", "Дата добавления", "Дата удаления"]
        extra_headers = [field.get("label", field.get("name")) for field in required_fields]
        headers = basic_headers + extra_headers

        cursor.close()

        # Выбор экспорта: Excel
        if export_option == "excel":
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(headers)
            column_widths = [40] * len(headers)
            for col_num, width in enumerate(column_widths, start=1):
                ws.column_dimensions[get_column_letter(col_num)].width = width
            for row in report_data:
                ws.append([str(cell) if cell is not None else "Нет данных" for cell in row])
            output = BytesIO()
            wb.save(output)
            output.seek(0)
            response = Response(
                output,
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            response.headers["Content-Disposition"] = f'attachment; filename="report_{access_base_id}.xlsx"'
            return response
        
        # Выбор экспорта: Word
        elif export_option == "word":
            output = generate_word_report(report_data, headers, base_name)
            response = Response(
                output,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response.headers["Content-Disposition"] = f'attachment; filename="report_{access_base_id}.docx"'
            return response
        
        else:
            flash("Отчёт успешно сформирован.", "success")
            return render_template("report.html", access_bases=[], data=report_data, headers=headers, selected_base=access_base_id)
    
    # GET-запрос — отображаем форму выбора базы доступа
    cursor.execute("SELECT AccessBaseID, AccessBaseName FROM AccessBases")
    access_bases = cursor.fetchall()
    cursor.close()
    return render_template("report.html", access_bases=access_bases, data=None, headers=None, selected_base=None)

if __name__ == "__main__":
    app.run(debug=True)
